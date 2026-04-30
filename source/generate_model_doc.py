"""
Generates a Word document that explains the bear-attack cost-benefit model
as it is currently implemented in source/pipeline.py.

Output: docs/bear_model_explained.docx
"""

from pathlib import Path
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

EXCEL_PATH = Path("source/bear_attack_model_variables.xlsx")
DOC_PATH = Path("docs/bear_model_explained.docx")


def read_variables(path: Path):
    wb = load_workbook(path)
    ws = wb.active
    rows = []
    for r in range(2, ws.max_row + 1):
        name = ws.cell(row=r, column=1).value
        if not name:
            continue
        rows.append({
            "name": name,
            "description": ws.cell(row=r, column=2).value,
            "notation": ws.cell(row=r, column=3).value,
            "value": ws.cell(row=r, column=4).value,
        })
    return rows


# ----- styling helpers -----

INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x4A, 0x6B, 0x3D)
DIM = RGBColor(0x55, 0x55, 0x55)


def set_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT if level == 1 else INK
    return h


def add_formula(doc, text):
    """Render a formula in monospace, indented, on its own line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.bold = True
    return p


def add_logic_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.color.rgb = DIM
    return p


def add_variable_table(doc, variables):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Notation"
    hdr[1].text = "Name"
    hdr[2].text = "Default"
    hdr[3].text = "Description"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for v in variables:
        row = table.add_row().cells
        notation = v["notation"] or ""
        # Fix mojibake for "not M": the Excel stored P(¬M_j) garbled
        notation = notation.replace("�", "¬").replace("\x00", "")
        row[0].text = notation
        row[1].text = v["name"] or ""
        val = v["value"]
        row[2].text = "" if val is None else (
            f"{val:.2f}" if isinstance(val, float) else str(val)
        )
        row[3].text = v["description"] or ""

    # Tighten font in the table
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)


# ----- document content -----

INTRODUCTION = (
    "This document describes the bear-attack cost-benefit model as it is "
    "currently implemented in the project pipeline (source/pipeline.py). "
    "The model compares two strategies for managing bear-caused damage on a "
    "single property over a fixed time horizon: paying compensation for "
    "every damage event (\"no protection\") versus installing and "
    "maintaining a protective measure on a share of the assets "
    "(\"with protection\")."
)

OVERVIEW = (
    "The pipeline reads the model variables from an Excel workbook, embeds "
    "them as defaults inside an interactive HTML page, and lets the user "
    "adjust each parameter through sliders. Whenever a slider moves, the "
    "front-end recomputes the formulas described below and redraws the "
    "per-element tables, the per-element cumulative-cost charts, and the "
    "overall verdict. The Excel workbook is therefore the single source of "
    "default values, and the formulas in this document are the exact "
    "expressions evaluated in the browser."
)

ASSETS_NOTE = (
    "The property is divided into three asset types, each indexed by i: "
    "small livestock (SL), large livestock (LL), and agriculture (Ag). "
    "Bear behaviour, exposure, compensation cost, and the protective "
    "measure are parameterised separately for each of them."
)

# Each entry: (heading, formula, logic paragraph)
FORMULAS = [
    (
        "1. Property-level attack rate λ (events per year)",
        "λ = P(B) · P(E|B) · P(D_i|food) · P(D_i|density) · Σ [ s_i · P(D_i | B, E) ]",
        "This expression collapses the original probability chain into a single "
        "rate of damage events per year on the property. P(B) gives the chance a "
        "bear is in the area at all; P(E|B) the chance such a bear actually "
        "encounters an asset in a way that could cause damage. The two "
        "environment factors — wild-food availability and bear-population "
        "density — multiplicatively scale that base rate up or down (values "
        "above 1 worsen pressure, below 1 ease it). The final sum averages the "
        "element-specific damage probabilities, weighted by how often each asset "
        "type is the target (the attack shares s_i). The result λ is the "
        "expected number of damage events on the property in one year, before "
        "splitting them between asset types."
    ),
    (
        "2. Element-level event rate",
        "λ_i = λ · s_i",
        "Once the property-level rate λ is known, it is divided among the three "
        "asset types according to their attack shares. s_SL + s_LL + s_Ag = 1 by "
        "construction (the front-end auto-normalises the shares whenever the "
        "user moves one slider), so this step simply allocates the total annual "
        "events without inventing or losing any. λ_i is what feeds every "
        "downstream cost calculation for asset type i."
    ),
    (
        "3. Annual loss without protection",
        "loss_i^noprot = λ_i · min(u_i, N_i) · c_i",
        "If no measure is in place, every event the bear initiates can reach the "
        "asset. Bears do not destroy a whole flock in one visit — they take a "
        "behavioural number of units u_i (3–5 sheep, 1 calf, several beehives). "
        "The min(u_i, N_i) cap reflects that on a small property the bear "
        "cannot damage more units than actually exist. Multiplying by the "
        "annual event rate λ_i and the per-unit compensation c_i converts the "
        "physical damage into expected euros lost per year."
    ),
    (
        "4a. Annual cost with protection — unprotected exposure",
        "unprot_i = λ_i · min(u_i, N_i · (1 − s_prot_i)) · c_i",
        "When a measure protects only a share s_prot_i of the units, the "
        "remaining (1 − s_prot_i) share is still fully exposed. On those "
        "unprotected units the bear behaves exactly as in the no-protection "
        "case, taking up to u_i units per event but limited by however many "
        "unprotected units actually exist. This term is the residual "
        "compensation cost that no measure can remove unless coverage reaches "
        "100%."
    ),
    (
        "4b. Annual cost with protection — residual losses on protected units",
        "residual_i = λ_i · P(¬M_j) · min(u_i, N_i · s_prot_i) · c_i",
        "Even on the protected share, the measure occasionally fails — a fence "
        "is left open, a battery dies, a guardian dog is absent. P(¬M_j) is "
        "that failure probability (auto-derived as 1 − P(M_j) in the front-end). "
        "When the measure has failed, the bear can damage protected units too, "
        "and the same behavioural cap u_i applies (now bounded by the number of "
        "protected units). A maintained measure is assumed to be fully "
        "effective, so this is the only way damage leaks through the protected "
        "side."
    ),
    (
        "4c. Annual cost with protection — maintenance",
        "maint_i / yr = c_maintenance_j · N_i · s_prot_i",
        "Keeping the measure functional has a recurring annual cost: replacing "
        "fence batteries, feeding guardian dogs, repainting beehive platforms. "
        "It scales linearly with the number of units actually under "
        "protection (N_i · s_prot_i), and uses a single per-unit maintenance "
        "rate c_maintenance_j that the user calibrates once for the chosen "
        "measure."
    ),
    (
        "4d. Annual cost with protection — installation (year 0 only)",
        "install_i = c_install_i · N_i · s_prot_i",
        "Installation is the up-front capital cost paid the moment the measure "
        "is deployed: building the fence, buying and training the dog, raising "
        "the beehive platforms. It is paid once, in year 0, and is "
        "element-specific because protecting a sheep costs much less than "
        "protecting a horse or a beehive. Because this cost does not recur, it "
        "appears as a one-time step in the cumulative curve, not as a slope."
    ),
    (
        "5. Cumulative cost over T years",
        "C_noprot(T) = T · Σ loss_i^noprot\n"
        "C_prot(T)   = Σ install_i + T · Σ ( unprot_i + residual_i + maint_i/yr )",
        "Aggregating across asset types and across years gives the two cost "
        "curves the user sees on the chart. Without protection the curve is a "
        "straight line through the origin: every year you accumulate the same "
        "expected compensation. With protection the curve starts at the "
        "installation cost and then grows on a (usually) gentler slope — "
        "unprotected exposure plus residual leakage plus maintenance. The year "
        "at which the protection curve first drops below the no-protection one "
        "is the payback year (the front-end labels it as the crossover)."
    ),
    (
        "6. Net benefit at the time horizon T",
        "Net(T) = C_noprot(T) − C_prot(T)",
        "The verdict in the header — \"Protect\" or \"Compensate\" — is just "
        "the sign of Net(T) at the chosen horizon. A positive value means the "
        "protective strategy has saved more in avoided compensation than it has "
        "cost in installation and maintenance over T years. A negative value "
        "means compensation is still cheaper. The front-end displays this as "
        "an absolute amount, with the crossover year as additional context "
        "when protection eventually pays off but has not yet over the chosen "
        "horizon."
    ),
]

CLOSING = (
    "Two implementation notes worth keeping in mind. First, the share sliders "
    "s_SL, s_LL, s_Ag are auto-normalised in the browser so that they always "
    "sum to one — moving one redistributes the remainder across the others "
    "proportionally. Second, the failure probability P(¬M_j) is not an "
    "independent parameter: it is recomputed from P(M_j) every time the user "
    "moves that slider, which is why the Excel sheet stores it for "
    "reference but it does not appear as its own control."
)


def main():
    print("Reading variables from Excel...")
    variables = read_variables(EXCEL_PATH)
    print(f"  loaded {len(variables)} variables")

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    set_default_style(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Bear-Attack Cost-Benefit Model")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Model specification as implemented in the pipeline")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = DIM

    add_heading(doc, "Introduction", level=1)
    doc.add_paragraph(INTRODUCTION)

    add_heading(doc, "How the pipeline uses the model", level=1)
    doc.add_paragraph(OVERVIEW)
    doc.add_paragraph(ASSETS_NOTE)

    add_heading(doc, "Variable dictionary", level=1)
    doc.add_paragraph(
        "The table below lists every model variable, its notation, its default "
        "value as stored in bear_attack_model_variables.xlsx, and a short "
        "description. Subscripts: i indexes asset type (SL = small livestock, "
        "LL = large livestock, Ag = agriculture); j indexes the protective "
        "measure."
    )
    add_variable_table(doc, variables)

    add_heading(doc, "Formulas and the logic behind each", level=1)
    doc.add_paragraph(
        "Each formula below is followed by a paragraph explaining the reasoning "
        "behind it. The variables match those in the dictionary above."
    )

    for heading, formula, logic in FORMULAS:
        add_heading(doc, heading, level=2)
        # support multi-line formulas
        for line in formula.split("\n"):
            add_formula(doc, line)
        add_logic_paragraph(doc, logic)

    add_heading(doc, "Implementation notes", level=1)
    doc.add_paragraph(CLOSING)

    doc.save(DOC_PATH)
    print(f"  wrote {DOC_PATH}")
    print("\nDocumentation generated.")


if __name__ == "__main__":
    main()
